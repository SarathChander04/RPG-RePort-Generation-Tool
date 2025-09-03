import os
from collections import defaultdict
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Configure this to match your Flask app's upload directory name
UPLOAD_BASE_DIR = "evidence_uploads"
SEVERITY_ORDER = ["Critical", "High", "Medium", "Low", "Informational"]
IMAGE_EXTS = {".png", ".jpg", ".jpeg"}

def add_heading_with_bookmark(doc, text, level, bookmark_id):
    paragraph = doc.add_heading(level=level)
    paragraph.add_run(text)
    tag = f"_Toc{bookmark_id}"
    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), str(bookmark_id))
    bm_start.set(qn("w:name"), tag)
    paragraph._p.insert(0, bm_start)
    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.append(bm_end)
    return paragraph

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run("" if text is None else str(text))
    run.bold = bold
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_table(doc, data, headers, style="Table Grid"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = style
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)
        if hdr[i].paragraphs and hdr[i].paragraphs[0].runs:
            hdr[i].paragraphs[0].runs[0].font.bold = True
    for row in data:
        cells = table.add_row().cells
        for i, item in enumerate(row):
            cells[i].text = "" if item is None else str(item)
    return table

def _build_class_summary(vulns, targets_by_vuln):
    class_counts = defaultdict(lambda: {s: 0 for s in SEVERITY_ORDER} | {"Total": 0})
    for vid, v in vulns:
        sev = v.get("severity") or "Low"
        for target in targets_by_vuln.get(vid, []):
            # For backward compatibility, allow dual or triple tuples
            if len(target) == 3:
                tgt, tclass, port = target
            else:
                tgt, tclass = target
                port = ""
            if tclass and str(tclass).strip():
                cls = str(tclass).strip()
                class_counts[cls][sev] += 1
                class_counts[cls]["Total"] += 1
    class_counts.pop("", None)
    return class_counts

def _add_class_summary_table(doc, class_counts):
    add_paragraph(doc, "Severity counts by Class", bold=True)
    headers = ["Class"] + SEVERITY_ORDER + ["Total"]
    if not class_counts:
        add_table(doc, [["—"] + ["0"] * len(SEVERITY_ORDER) + ["0"]], headers)
        return
    rows = sorted(class_counts.items(), key=lambda kv: (-kv[1]["Total"], kv[0].lower()))
    data = []
    for cls, counts in rows:
        data.append([cls] + [counts.get(s, 0) for s in SEVERITY_ORDER] + [counts.get("Total", 0)])
    add_table(doc, data, headers)

def _add_targets_table(doc, targets):
    """
    Three-column table: Target | Class | Port
    """
    headers = ["Target", "Class", "Port"]
    if not targets:
        add_table(doc, [["No targets", "-", "-"]], headers)
        return
    data = []
    for target in targets:
        if len(target) == 3:
            tgt, tclass, port = target
        else:
            tgt, tclass = target
            port = ""
        data.append([tgt or "", tclass or "", port or ""])
    add_table(doc, data, headers)

def _is_image_path(path_str):
    try:
        ext = os.path.splitext(path_str)[1].lower()
        return ext in IMAGE_EXTS
    except Exception:
        return False

def _resolve_evidence_path(evidence_value):
    if not evidence_value:
        return None, ""
    ev = str(evidence_value).strip()
    if os.path.isabs(ev):
        return (ev if os.path.exists(ev) else None), ev
    # Relative under upload base
    candidate = os.path.join(UPLOAD_BASE_DIR, ev)
    if os.path.exists(candidate):
        return candidate, ev
    candidate2 = os.path.abspath(ev)
    if os.path.exists(candidate2):
        return candidate2, ev
    return None, ev

def _add_evidence(doc, evidence_value):
    add_paragraph(doc, "Evidence", bold=True)
    if not evidence_value:
        add_paragraph(doc, "No evidence provided.")
        return
    abs_path, display_text = _resolve_evidence_path(evidence_value)
    if abs_path and _is_image_path(abs_path):
        try:
            doc.add_picture(abs_path, width=Inches(5.5))
            add_paragraph(doc, os.path.basename(abs_path))
            return
        except Exception as e:
            add_paragraph(doc, f"[Could not embed image: {e}]")
    add_paragraph(doc, display_text or str(evidence_value))

def generate_word_report(company, vulns, targets_by_vuln, output_path):
    document = Document()

    # Title
    title = document.add_heading("Vulnerability Assessment and Penetration Testing Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(document, f"Company: {company}", bold=True)
    document.add_page_break()

    # Confidentiality
    document.add_heading("Confidentiality Notice", level=1)
    add_paragraph(
        document,
        "This Report is prepared solely for the designated company and contains information that should be considered as Confidential.\n\n"
        "The use may not meet the requirements or objectives of any other party. This report is confidential. It includes "
        "technical details of security weaknesses that if available to untrusted parties could facilitate the execution of a breach "
        "against the client. For that reason, access to this report should be restricted, and the report should not be published."
    )
    document.add_page_break()

    # Executive Summary with Class-wise summary
    document.add_heading("Executive Summary", level=1)
    add_paragraph(
        document,
        "Company engaged with Alliance Pro to perform a Technical Vulnerability Assessment and Penetration Testing on their web applications.\n\n"
        "Our testing approach combines the use of automated scanning tools and manual techniques on the web application. "
        "This assists with the identification of logical security vulnerabilities, patching deficiencies, and misconfigurations, "
        "to make mitigation recommendations."
    )
    add_paragraph(
        document,
        "Our work is based at a specific point in time, in an environment where both the systems and the threat profiles are dynamically "
        "evolving. It is therefore possible that vulnerabilities exist or will arise that were not identified during our review and "
        "there may or will have been events, developments, and changes in circumstances after its issue which is likely to render the "
        "report’s contents wholly or partly misleading or inaccurate."
    )

    document.add_heading("Scope of a Web Application Penetration Testing and the severity wise vulnerabilities.", level=2)
    add_paragraph(document, "Scope covers in-scope web applications, network-connected systems, APIs, and their related infrastructure, as provided by the client.")
    add_paragraph(document, "The assessment's duration and methodology are described in later sections. Only security issues detected at the time of testing are listed; the testing is non-exhaustive.")
    class_counts = _build_class_summary(vulns, targets_by_vuln)
    _add_class_summary_table(document, class_counts)
    add_paragraph(document, "")
    add_paragraph(
        document,
        "This assessment presents the findings and gives information on any potential vulnerabilities in the Web Application. "
        "It is recommended that additional security measures be put in place to protect the confidentiality of both this document and "
        "the data it contains. The recommendations in this report are founded on the information that is currently available from the "
        "scan results and manual penetration testing. Vulnerability Assessment and Penetration Testing is one method for Web "
        "Applications; among others, performing a source code review would be another method."
    )
    add_paragraph(
        document,
        "The Common Vulnerabilities and Exposures (CVE) system, which serves as a reference method for publicly known information-security "
        "vulnerabilities and exposures, is used to provide remediation advice. The development team can implement the recommendations across "
        "the web application to close significant vulnerabilities."
    )
    document.add_page_break()

    # Summary of Vulnerabilities Discovered
    document.add_heading("Summary of Vulnerabilities Discovered", level=1)
    add_paragraph(
        document,
        "The table below summarizes all vulnerabilities discovered during this assessment, including their severity and the total number of affected URLs/paths."
    )
    summary_data = [
        [v.get("vuln_name", ""), v.get("severity", ""), len(targets_by_vuln.get(vid, []))]
        for vid, v in vulns
    ]
    add_table(document, summary_data, ["Vulnerability Name", "Severity", "Total No. of Affected URLs/Paths"])
    document.add_page_break()

    # Methodology
    document.add_heading("Our Methodology", level=1)
    add_paragraph(
        document,
        "We used a methodology based on the globally accepted standards and practices. Alliance independently tested the identified components using non-intrusive methods, to evaluate the existing security measures pertaining to the system. The security issues observed are presented in this report together with an assessment of their likely impact. Our team has followed a Black Box testing approach for the assessment. This approach simulates the activities that of an attacker or disgruntled employee."
    )
    add_paragraph(
        document,
        "Summary of the process:"
    )
    add_paragraph(
        document,
        "Phase 1: Reconnaissance\n"
        "During this phase, the scope of the testing was identified using the information available without any restrictions, using scanning and crawling tools. Using this information, web application page structure, active systems and services/applications running on the system were mapped."
    )
    add_paragraph(
        document,
        "Phase 2: Vulnerability Scanning\n"
        "The discovered services and applications were compared with the vulnerability database and outlined vulnerabilities using automated vulnerability scanners. In this process, detailed information about the host operating systems, web application components, and any filtered services running on these hosts was obtained."
    )
    add_paragraph(
        document,
        "Phase 3: Vulnerability Analysis\n"
        "After understanding the vulnerabilities applicable to environment, the team tried to assess the impact of the vulnerabilities on the identified application(s) and system(s) and categorized them based on risk levels."
    )
    add_paragraph(
        document,
        "Phase 4: Penetration Testing\n"
        "The identified vulnerabilities in the environment were then exploited as a part of the penetration testing process in a controlled manner, with adequate proof provided to showcase the imminent impact of successful exploitation of these vulnerabilities."
    )
    add_paragraph(
        document,
        "Phase 5: Reporting\n"
        "A final document containing the information and statistics about the results of the vulnerability assessment exercise was prepared, along with detailed information and recommendations as per the industry best practices."
    )
    document.add_page_break()

    # Key Issues
    document.add_heading("Key Issues", level=1)
    add_paragraph(
        document,
        "Overall, it was seen that while having established fundamental rules for safeguarding the web application perimeter, there are substantial weaknesses present on the web applications and all the internal systems are completely secured with all ports being filtered. It's important to fix these flaws within a reasonable amount of time. As a result of only one high vulnerability found and the numerous medium-risk vulnerabilities found during this analysis, it might be possible that using outdated components and versions may lead to exploitation of publicly available exploits."
    )
    document.add_page_break()

    # Risk Rating Definitions
    document.add_heading("Risk Rating Definitions", level=1)
    add_paragraph(
        document,
        "Observations made during this assessment have been identified with one of the following risk levels. Each risk level indicates the significance and likelihood of specific risk types in the IT environment. This assessment can be used by management as a tool to determine how quickly attention should be given to each observation provided within this report."
    )
    risk_rating_data = [
        ["Critical", "A critical-risk vulnerability can be exploited to execute code or access sensitive data, often with little knowledge required."],
        ["High", "A high-risk vulnerability can be exploited with ease and little or no authentication, potentially exposing confidential information or causing outages."],
        ["Medium", "A medium-risk vulnerability may require moderate expertise or authentication and can allow partial access or limited disruption."],
        ["Low", "A low-risk vulnerability typically requires local access and authentication and has limited impact on confidentiality, integrity, and availability."],
        ["Informational", "Issues that are not vulnerabilities but may assist attackers or indicate potential weaknesses if combined with other issues."]
    ]
    add_table(document, risk_rating_data, ["Risk Level", "Description"])
    add_paragraph(document, "Note: These definitions may vary according to specific industry verticals.")
    document.add_page_break()

    # Detailed Vulnerability Information (ordered by severity then name)
    document.add_heading("Detailed Vulnerability Information", level=1)
    add_paragraph(
        document,
        "This section provides in-depth technical details for each finding, along with affected targets/URLs, the evidence provided, and remediation recommendations."
    )
    sev_rank = {s: i for i, s in enumerate(SEVERITY_ORDER)}
    ordered = sorted(
        vulns,
        key=lambda kv: (sev_rank.get((kv[1].get("severity") or "Low"), 5), kv[1].get("vuln_name", "").lower())
    )

    bookmark_counter = 1000
    for idx, (vuln_id, v) in enumerate(ordered, 1):
        add_heading_with_bookmark(document, f"{idx}. {v.get('vuln_name','')}", level=2, bookmark_id=bookmark_counter)
        bookmark_counter += 1

        add_paragraph(document, f"Severity: {v.get('severity','')}")
        if v.get("cvss_score") is not None:
            add_paragraph(document, f"CVSS Score: {v.get('cvss_score','')}")
        add_paragraph(document, f"Description: {v.get('description','')}")
        add_paragraph(document, f"Impact: {v.get('impact','')}")
        add_paragraph(document, f"Remediation: {v.get('remediation','')}")
        add_paragraph(document, "Affected Targets, Classes, and Ports", bold=True)
        _add_targets_table(document, targets_by_vuln.get(vuln_id, []))
        _add_evidence(document, v.get("evidence"))
        if idx < len(ordered):
            document.add_page_break()
    document.add_page_break()

    # General Recommendations
    document.add_heading("General Recommendations", level=1)
    add_paragraph(
        document,
        "Based on the assessment, the following actions are recommended to improve the web application's security posture:"
    )
    recommended = (
        "• **Secure Configuration:** Implement secure configurations for all components of the web application stack, including minimizing unnecessary services, disabling default accounts, and utilizing strong encryption protocols.\n"
        "• **Input Validation:** Implement rigorous input validation techniques to prevent injection attacks such as SQL injection, cross-site scripting (XSS), and command injection. Validate and sanitize all user inputs before processing.\n"
        "• **Output Encoding:** Encode output to prevent XSS attacks. Use output encoding libraries appropriate for the output context (HTML, URL, JavaScript, etc.).\n"
        "• **Session Management:** Implement secure session management mechanisms, including unique session IDs, session expiration, secure cookie attributes, and protection against session fixation attacks.\n"
        "• **Error Handling:** Implement proper error handling to avoid leaking sensitive information and provide minimal information to users in error messages. Log errors securely for debugging purposes without exposing sensitive information.\n"
        "• **Data Encryption:** Encrypt sensitive data at rest and in transit using strong cryptographic algorithms. Utilize HTTPS with secure TLS configurations to protect data in transit.\n"
        "• **Security Headers:** Implement appropriate security headers such as Content Security Policy (CSP), Strict Transport Security (HSTS), and X-Frame-Options to enhance the security posture of the application.\n"
        "• **Regular Updates and Patching:** Keep all software components, including the operating system, web server, frameworks, and libraries, up to date with the latest security patches to mitigate known vulnerabilities.\n"
        "• **Source Code Review:** Conduct regular manual and automated source code reviews to identify and fix security flaws at the code level.\n"
        "• **Security Training:** Provide regular security awareness training to developers and other stakeholders to ensure they are aware of the latest threats and best practices for secure coding.\n"
        "• **Web Application Firewall (WAF):** Deploy a WAF to provide an additional layer of protection against common web attacks by filtering malicious traffic."
    )
    for line in recommended.split("\n"):
        add_paragraph(document, line.lstrip("• "), bold=line.startswith("• "))

    document.save(output_path)
    return True
